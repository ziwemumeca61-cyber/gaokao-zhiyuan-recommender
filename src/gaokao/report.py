"""一键导出推荐理由文档。

把考生画像 + 冲稳保志愿表 + 每条志愿的概率/兴趣/综合分/推荐理由整理成文档；
也支持把考生心愿单（已选 院校·专业）导出成一张"志愿表"。
默认输出 Markdown（零依赖）；可选输出 Word（python-docx）/PDF（reportlab），
依赖缺失时自动不可用，由调用方判断。
"""

from __future__ import annotations

from datetime import date

from .models import TIERS, Major, RIASEC_LABELS, Recommendation, School, Student

_DISCLAIMER = "仅供参考，请结合各校招生章程谨慎决策"


def _scope_note(student: Student) -> str:
    """省份口径说明：报告内所有分数线/位次均为该省该科类，跨省不可比。"""
    return (f"数据口径：{student.province}·{student.subject_type}"
            "（高考按省份分别划线录取，跨省分数/位次不可直接比较）")


def _intro(major: Major) -> str:
    """专业简介：取自带或精选库；未收录则给诚实提示，不用门类模板硬凑。"""
    from .major_knowledge import detail_for  # noqa: PLC0415

    d = detail_for(major)
    if d["covered"]:
        return d["intro"]
    return "暂未收录该专业的详细介绍，建议查阅目标院校培养方案。"


def _advice(major: Major) -> dict[str, str]:
    """专业选报建议（选专业提示/实在话/适合谁）。"""
    from .major_advice import advice_for  # noqa: PLC0415

    return advice_for(major.name, major.category)


# ---------------------------------------------------------------------------
# 考生画像（多种格式共用的文本片段）
# ---------------------------------------------------------------------------
def _profile_lines(student: Student) -> list[str]:
    lines = [
        f"分数 {student.score}　位次 {student.rank}",
        f"{student.province}　{student.subject_type}",
    ]
    if student.level_pref:
        lines.append(f"院校层次偏好：{student.level_pref}")
    if student.city_prefs:
        lines.append(f"城市偏好：{', '.join(student.city_prefs)}")
    if student.major_prefs:
        lines.append(f"专业门类偏好：{', '.join(student.major_prefs)}")
    if student.has_assessment():
        profile = "　".join(
            f"{RIASEC_LABELS[d]} {v:.2f}" for d, v in student.riasec.items())
        lines.append(f"兴趣画像(RIASEC)：{profile}")
    return lines


# ===========================================================================
# 一、冲稳保推荐报告
# ===========================================================================
def build_markdown_report(
    student: Student, buckets: dict[str, list[Recommendation]]
) -> str:
    lines: list[str] = []
    lines.append("# 高考志愿推荐报告")
    lines.append("")
    lines.append(f"> 生成日期：{date.today().isoformat()}　|　{_DISCLAIMER}")
    lines.append(">")
    lines.append(f"> 📍 {_scope_note(student)}")
    lines.append("")

    # 考生信息
    lines.append("## 一、考生画像")
    lines.append("")
    for ln in _profile_lines(student):
        lines.append(f"- {ln}")
    lines.append("")

    # 整体策略
    counts = {t: len(buckets.get(t, [])) for t in TIERS}
    lines.append("## 二、整体策略")
    lines.append("")
    lines.append(f"- 冲 {counts['冲']} 个 / 稳 {counts['稳']} 个 / 保 {counts['保']} 个")
    lines.append("- 建议按'冲稳保'梯度合理搭配，保底志愿务必稳妥，避免滑档。")
    lines.append("")

    # 各档志愿
    lines.append("## 三、志愿推荐明细")
    for tier in TIERS:
        recs = buckets.get(tier, [])
        lines.append("")
        lines.append(f"### {tier}（{len(recs)} 个）")
        if not recs:
            lines.append("")
            lines.append("（暂无）")
            continue
        for i, rec in enumerate(recs, start=1):
            lines.append("")
            lines.append(
                f"**{i}. {rec.school.name} · {rec.major.name}**"
                f"（{rec.school.level}/{rec.school.city}）")
            lines.append(
                f"- 参考位次 {rec.ref_rank}　参考分 {rec.ref_score}　"
                f"录取概率 {rec.probability * 100:.0f}%"
                f"（{rec.prob_low * 100:.0f}–{rec.prob_high * 100:.0f}%，把握度{rec.confidence}）"
                f"　综合分 {rec.composite_score:.2f}")
            lines.append(f"- 专业简介：{_intro(rec.major)}")
            if rec.reasons:
                lines.append(f"- 推荐理由：{'；'.join(rec.reasons)}")
    lines.append("")
    return "\n".join(lines)


# ===========================================================================
# 二、心愿单（志愿表）
# ===========================================================================
def build_markdown_wishlist(
    student: Student, items: list[tuple[School | None, Major]]
) -> str:
    """把心愿单（按填报顺序的 院校·专业）导出为 Markdown 志愿表。"""
    lines: list[str] = []
    lines.append("# 我的志愿表")
    lines.append("")
    lines.append(f"> 生成日期：{date.today().isoformat()}　|　{_DISCLAIMER}")
    lines.append(">")
    lines.append(f"> 📍 {_scope_note(student)}")
    lines.append("")
    lines.append("## 考生画像")
    lines.append("")
    for ln in _profile_lines(student):
        lines.append(f"- {ln}")
    lines.append("")

    if not items:
        lines.append("## 志愿顺序（共 0 个）")
        lines.append("")
        lines.append("（心愿单为空，先去添加心仪的专业吧）")
        return "\n".join(lines)

    # 志愿诊断（整体体检）
    from .diagnosis import diagnose  # noqa: PLC0415

    diag = diagnose(student, items)
    lines.append("## 志愿诊断")
    lines.append("")
    for _sev, text in diag.findings:
        lines.append(f"- {text}")
    lines.append("")

    lines.append(f"## 志愿顺序（共 {len(items)} 个）")

    for i, (school, major) in enumerate(items, start=1):
        loc = f"（{school.level}/{school.city}）" if school else ""
        sname = school.name if school else "未知院校"
        lines.append("")
        lines.append(f"**{i}. {sname} · {major.name}**{loc}")
        lines.append(f"- 学科门类 {major.category}")
        lines.append(f"- 专业简介：{_intro(major)}")
        _a = _advice(major)
        lines.append(f"- ⚠️ 选专业提示：{_a['pitfall']}")
        lines.append(f"- 💬 实在话：{_a['truth']}")
        lines.append(f"- 👪 适合谁：{_a['fit']}")
    lines.append("")
    lines.append("> 注：以上专业点评为行业普遍看法，仅供参考。")
    return "\n".join(lines)


# ===========================================================================
# 零依赖兜底：Markdown -> RTF（Word 可开）/ 可打印 HTML（浏览器另存为 PDF）
# 无需 python-docx / reportlab，保证任何环境都能导出。
# ===========================================================================
def _strip_bold(s: str) -> str:
    return s.replace("**", "")


def _md_lines(md: str):
    """把我们生成的简单 Markdown 解析成 (类型, 文本) 序列。"""
    for raw in md.split("\n"):
        line = raw.rstrip()
        if not line:
            yield ("blank", "")
        elif line.startswith("# "):
            yield ("h1", line[2:])
        elif line.startswith("## "):
            yield ("h2", line[3:])
        elif line.startswith("> "):
            yield ("quote", line[2:])
        elif line.startswith("- "):
            yield ("li", line[2:])
        elif line == ">":
            yield ("blank", "")
        else:
            yield ("p", line)


def _html_escape(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _md_to_html(md: str, title: str) -> bytes:
    """转为自带样式、适合打印为 PDF 的 HTML。"""
    body: list[str] = []
    in_ul = False

    def close_ul():
        nonlocal in_ul
        if in_ul:
            body.append("</ul>")
            in_ul = False

    for kind, text in _md_lines(md):
        # 行内加粗 **x** -> <b>
        html = _html_escape(text)
        while "**" in html:
            html = html.replace("**", "<b>", 1).replace("**", "</b>", 1)
        if kind == "li":
            if not in_ul:
                body.append("<ul>")
                in_ul = True
            body.append(f"<li>{html}</li>")
            continue
        close_ul()
        if kind == "h1":
            body.append(f"<h1>{html}</h1>")
        elif kind == "h2":
            body.append(f"<h2>{html}</h2>")
        elif kind == "quote":
            body.append(f"<p class='note'>{html}</p>")
        elif kind == "p":
            body.append(f"<p>{html}</p>")
    close_ul()
    doc = (
        "<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'>"
        f"<title>{_html_escape(title)}</title><style>"
        "body{font-family:'Microsoft YaHei','PingFang SC',sans-serif;max-width:780px;"
        "margin:24px auto;padding:0 16px;color:#222;line-height:1.7}"
        "h1{font-size:24px}h2{font-size:18px;border-bottom:1px solid #ddd;padding-bottom:4px;margin-top:24px}"
        ".note{color:#888;font-size:13px}.tip{background:#EAF6EE;border:1px solid #1FA463;"
        "padding:8px 12px;border-radius:6px;color:#1a7a47;font-size:13px}"
        "li{margin:2px 0}@media print{.tip{display:none}}</style></head><body>"
        "<p class='tip'>💡 用浏览器打开本文件后，按 Ctrl/⌘ + P → 目标选择“另存为 PDF”即可导出 PDF。</p>"
        + "".join(body) + "</body></html>"
    )
    return doc.encode("utf-8")


def _rtf_text(s: str) -> str:
    """RTF 文本转义：反斜杠/花括号转义，非 ASCII 用 \\uN? 表示（含中文）。"""
    out: list[str] = []
    for ch in s:
        if ch in "\\{}":
            out.append("\\" + ch)
        elif ord(ch) > 127:
            cp = ord(ch)
            if cp > 32767:
                cp -= 65536
            out.append(f"\\u{cp}?")
        else:
            out.append(ch)
    return "".join(out)


def _md_to_rtf(md: str) -> bytes:
    """转为 RTF（Word/WPS 可直接打开），中文用 Unicode 转义，无需任何库。"""
    parts: list[str] = [r"{\rtf1\ansi\ansicpg936\deff0"
                        r"{\fonttbl{\f0\fnil\fcharset134 SimSun;}}\f0\fs22"]
    for kind, text in _md_lines(md):
        if kind == "blank":
            parts.append(r"\par")
            continue
        body = _rtf_text(_strip_bold(text))
        if kind == "h1":
            parts.append(r"\pard\sa120\b\fs36 " + body + r"\b0\par")
        elif kind == "h2":
            parts.append(r"\pard\sa80\b\fs28 " + body + r"\b0\par")
        elif kind == "li":
            parts.append(r"\pard\fi-280\li360\sa40 \bullet  " + body + r"\par")
        elif kind == "quote":
            parts.append(r"\pard\sa40\i\fs18 " + body + r"\i0\par")
        else:
            parts.append(r"\pard\sa40 " + body + r"\par")
    parts.append("}")
    return "".join(parts).encode("utf-8")


def build_rtf_report(student: Student, buckets: dict[str, list[Recommendation]]) -> bytes:
    return _md_to_rtf(build_markdown_report(student, buckets))


def build_rtf_wishlist(student: Student, items: list[tuple[School | None, Major]]) -> bytes:
    return _md_to_rtf(build_markdown_wishlist(student, items))


def build_html_report(student: Student, buckets: dict[str, list[Recommendation]]) -> bytes:
    return _md_to_html(build_markdown_report(student, buckets), "高考志愿推荐报告")


def build_html_wishlist(student: Student, items: list[tuple[School | None, Major]]) -> bytes:
    return _md_to_html(build_markdown_wishlist(student, items), "我的志愿表")


# ===========================================================================
# Word（python-docx）
# ===========================================================================
def docx_available() -> bool:
    try:
        import docx  # noqa: F401, PLC0415

        return True
    except Exception:
        return False


def build_docx_report(student: Student, buckets: dict[str, list[Recommendation]]) -> bytes:
    """生成 Word 文档字节流（需要 python-docx）。"""
    from docx import Document  # noqa: PLC0415

    doc = Document()
    doc.add_heading("高考志愿推荐报告", level=0)
    doc.add_paragraph(f"生成日期：{date.today().isoformat()}（{_DISCLAIMER}）")
    doc.add_paragraph(f"📍 {_scope_note(student)}")

    doc.add_heading("考生画像", level=1)
    for ln in _profile_lines(student):
        doc.add_paragraph(ln, style="List Bullet")

    for tier in TIERS:
        recs = buckets.get(tier, [])
        doc.add_heading(f"{tier}（{len(recs)} 个）", level=1)
        for i, rec in enumerate(recs, start=1):
            doc.add_heading(f"{i}. {rec.school.name} · {rec.major.name}", level=2)
            doc.add_paragraph(
                f"{rec.school.level}/{rec.school.city}　参考位次 {rec.ref_rank}　"
                f"录取概率 {rec.probability * 100:.0f}%"
                f"（{rec.prob_low * 100:.0f}–{rec.prob_high * 100:.0f}%，把握度{rec.confidence}）")
            doc.add_paragraph(f"专业简介：{_intro(rec.major)}")
            if rec.reasons:
                doc.add_paragraph("推荐理由：" + "；".join(rec.reasons))

    return _docx_bytes(doc)


def build_docx_wishlist(
    student: Student, items: list[tuple[School | None, Major]]
) -> bytes:
    """生成心愿单（志愿表）的 Word 字节流。"""
    from docx import Document  # noqa: PLC0415

    doc = Document()
    doc.add_heading("我的志愿表", level=0)
    doc.add_paragraph(f"生成日期：{date.today().isoformat()}（{_DISCLAIMER}）")
    doc.add_paragraph(f"📍 {_scope_note(student)}")

    doc.add_heading("考生画像", level=1)
    for ln in _profile_lines(student):
        doc.add_paragraph(ln, style="List Bullet")

    if items:
        from .diagnosis import diagnose  # noqa: PLC0415

        doc.add_heading("志愿诊断", level=1)
        for _sev, text in diagnose(student, items).findings:
            doc.add_paragraph(text, style="List Bullet")

    doc.add_heading(f"志愿顺序（共 {len(items)} 个）", level=1)
    for i, (school, major) in enumerate(items, start=1):
        sname = school.name if school else "未知院校"
        loc = f"（{school.level}/{school.city}）" if school else ""
        doc.add_heading(f"{i}. {sname} · {major.name}", level=2)
        doc.add_paragraph(f"{loc}学科门类 {major.category}")
        doc.add_paragraph(f"专业简介：{_intro(major)}")
        _a = _advice(major)
        doc.add_paragraph(f"⚠️ 选专业提示：{_a['pitfall']}")
        doc.add_paragraph(f"💬 实在话：{_a['truth']}")
        doc.add_paragraph(f"👪 适合谁：{_a['fit']}")

    doc.add_paragraph("注：以上专业点评为行业普遍看法，仅供参考。")
    return _docx_bytes(doc)


def _docx_bytes(doc) -> bytes:
    import io  # noqa: PLC0415

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===========================================================================
# PDF（reportlab）—— 用内置 CID 字体 STSong-Light 直接支持中文，无需外部字体
# ===========================================================================
def pdf_available() -> bool:
    try:
        import reportlab  # noqa: F401, PLC0415

        return True
    except Exception:
        return False


def build_pdf_report(student: Student, buckets: dict[str, list[Recommendation]]) -> bytes:
    """生成冲稳保推荐报告的 PDF 字节流（需要 reportlab）。"""
    blocks: list[tuple[str, str]] = [
        ("高考志愿推荐报告", "title"),
        (f"生成日期：{date.today().isoformat()}（{_DISCLAIMER}）", "caption"),
        (_scope_note(student), "caption"),
        ("一、考生画像", "h1"),
    ]
    blocks += [(ln, "body") for ln in _profile_lines(student)]

    counts = {t: len(buckets.get(t, [])) for t in TIERS}
    blocks.append(("二、整体策略", "h1"))
    blocks.append((f"冲 {counts['冲']} 个 / 稳 {counts['稳']} 个 / 保 {counts['保']} 个", "body"))
    blocks.append(("建议按'冲稳保'梯度合理搭配，保底志愿务必稳妥，避免滑档。", "body"))

    blocks.append(("三、志愿推荐明细", "h1"))
    for tier in TIERS:
        recs = buckets.get(tier, [])
        blocks.append((f"{tier}（{len(recs)} 个）", "h2"))
        for i, rec in enumerate(recs, start=1):
            blocks.append((
                f"{i}. {rec.school.name} · {rec.major.name}"
                f"（{rec.school.level}/{rec.school.city}）", "h3"))
            blocks.append((
                f"参考位次 {rec.ref_rank}　参考分 {rec.ref_score}　"
                f"录取概率 {rec.probability * 100:.0f}%"
                f"（{rec.prob_low * 100:.0f}–{rec.prob_high * 100:.0f}%，把握度{rec.confidence}）"
                f"　综合分 {rec.composite_score:.2f}", "body"))
            if rec.reasons:
                blocks.append(("推荐理由：" + "；".join(rec.reasons), "body"))
    return _pdf_bytes(blocks)


def build_pdf_wishlist(
    student: Student, items: list[tuple[School | None, Major]]
) -> bytes:
    """生成心愿单（志愿表）的 PDF 字节流。"""
    blocks: list[tuple[str, str]] = [
        ("我的志愿表", "title"),
        (f"生成日期：{date.today().isoformat()}（{_DISCLAIMER}）", "caption"),
        (_scope_note(student), "caption"),
        ("考生画像", "h1"),
    ]
    blocks += [(ln, "body") for ln in _profile_lines(student)]
    if items:
        from .diagnosis import diagnose  # noqa: PLC0415

        blocks.append(("志愿诊断", "h1"))
        blocks += [(text, "body") for _sev, text in diagnose(student, items).findings]
    blocks.append((f"志愿顺序（共 {len(items)} 个）", "h1"))
    for i, (school, major) in enumerate(items, start=1):
        sname = school.name if school else "未知院校"
        loc = f"（{school.level}/{school.city}）" if school else ""
        blocks.append((f"{i}. {sname} · {major.name}{loc}", "h3"))
        blocks.append((f"学科门类 {major.category}", "body"))
        blocks.append(("专业简介：" + _intro(major), "body"))
        _a = _advice(major)
        blocks.append(("选专业提示：" + _a["pitfall"], "body"))
        blocks.append(("实在话：" + _a["truth"], "body"))
        blocks.append(("适合谁：" + _a["fit"], "body"))
    blocks.append(("注：以上专业点评为行业普遍看法，仅供参考。", "caption"))
    return _pdf_bytes(blocks)


def _esc(text: str) -> str:
    """转义 reportlab Paragraph 的 XML 特殊字符。"""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_bytes(blocks: list[tuple[str, str]]) -> bytes:
    """把 (文本, 样式) 块列表渲染为 PDF 字节流。样式：title/caption/h1/h2/h3/body。"""
    import io  # noqa: PLC0415

    from reportlab.lib.enums import TA_CENTER  # noqa: PLC0415
    from reportlab.lib.pagesizes import A4  # noqa: PLC0415
    from reportlab.lib.styles import ParagraphStyle  # noqa: PLC0415
    from reportlab.lib.units import mm  # noqa: PLC0415
    from reportlab.pdfbase import pdfmetrics  # noqa: PLC0415
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # noqa: PLC0415
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # noqa: PLC0415

    font = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font))
    except Exception:
        font = "Helvetica"  # 极端兜底（理论上 CID 字体随 reportlab 自带）

    styles = {
        "title": ParagraphStyle("t", fontName=font, fontSize=20, leading=26,
                                 alignment=TA_CENTER, spaceAfter=6),
        "caption": ParagraphStyle("c", fontName=font, fontSize=9, leading=13,
                                   alignment=TA_CENTER, textColor="#888888", spaceAfter=10),
        "h1": ParagraphStyle("h1", fontName=font, fontSize=15, leading=20,
                             spaceBefore=10, spaceAfter=4),
        "h2": ParagraphStyle("h2", fontName=font, fontSize=13, leading=18,
                             spaceBefore=8, spaceAfter=3),
        "h3": ParagraphStyle("h3", fontName=font, fontSize=11, leading=16,
                             spaceBefore=6, spaceAfter=2),
        "body": ParagraphStyle("b", fontName=font, fontSize=10, leading=15),
    }

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm,
                            leftMargin=18 * mm, rightMargin=18 * mm, title="志愿报告")
    flow: list = []
    for text, kind in blocks:
        flow.append(Paragraph(_esc(text), styles.get(kind, styles["body"])))
        if kind == "title":
            flow.append(Spacer(1, 2))
    doc.build(flow)
    return buf.getvalue()
